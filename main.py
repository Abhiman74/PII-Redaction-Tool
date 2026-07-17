#!/usr/bin/env python3
import os
import sys
import logging
import argparse

# Setup logging configuration
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
    handlers=[
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger("main")

# Add the project directory to path
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from src.redactor import DocxRedactor

def setup_directories():
    """Ensure all required input/output directories exist."""
    os.makedirs("input", exist_ok=True)
    os.makedirs("output", exist_ok=True)
    os.makedirs("evaluation", exist_ok=True)

def convert_pdf_to_docx(pdf_path: str, docx_path: str):
    """Convert the PDF prospectus to a DOCX file using pdf2docx."""
    logger.info("Converting PDF '%s' to DOCX '%s'...", pdf_path, docx_path)
    try:
        from pdf2docx import Converter
        cv = Converter(pdf_path)
        # Convert all pages (start=0, end=None)
        cv.convert(docx_path, start=0, end=None)
        cv.close()
        logger.info("Successfully converted PDF to DOCX.")
    except ImportError as e:
        logger.error("pdf2docx is not installed: %s", e)
        raise RuntimeError("pdf2docx is required but not installed.") from e
    except Exception as e:
        logger.error("Failed to convert PDF to DOCX: %s", e)
        raise IOError(f"Failed to convert PDF to DOCX: {e}") from e

def generate_evaluation_template(input_docx: str, template_path: str):
    """
    Scan the document and generate a CSV template containing predictions
    that can be marked by a reviewer for NLP metrics evaluation.
    """
    logger.info("Generating ground-truth annotation template...")
    try:
        import pandas as pd
        from docx import Document
        from src.detector import HybridDetector
        
        doc = Document(input_docx)
        detector = HybridDetector()
        
        records = []
        p_index = 0
        
        # Helper to process paragraphs and collect entities
        def check_text(text: str, location_type: str, item_id: int):
            if not text.strip():
                return
            entities = detector.detect(text)
            for ent in entities:
                records.append({
                    "Location Type": location_type,
                    "Item ID": item_id,
                    "Original Text": ent.text,
                    "Entity Type": ent.entity_type,
                    "Start Offset": ent.start,
                    "End Offset": ent.end,
                    "Review (TP/FP/FN)": "TP"  # Pre-fill with TP as prediction, reviewer verifies
                })

        # Body paragraphs
        for p in doc.paragraphs:
            check_text(p.text, "Body Paragraph", p_index)
            p_index += 1
            
        # Tables
        t_index = 0
        for table in doc.tables:
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    for p_idx, p in enumerate(cell.paragraphs):
                        check_text(p.text, f"Table_{t_index}_Row_{r_idx}_Col_{c_idx}_Para_{p_idx}", p_idx)
            t_index += 1

        # Headers and footers
        s_index = 0
        for sec in doc.sections:
            if sec.header:
                for idx, p in enumerate(sec.header.paragraphs):
                    check_text(p.text, f"Section_{s_index}_Header_Para_{idx}", idx)
            if sec.footer:
                for idx, p in enumerate(sec.footer.paragraphs):
                    check_text(p.text, f"Section_{s_index}_Footer_Para_{idx}", idx)
            s_index += 1

        df = pd.DataFrame(records)
        # Drop duplicates to keep template clean
        if not df.empty:
            df = df.drop_duplicates(subset=["Location Type", "Original Text", "Entity Type", "Start Offset"])
        df.to_csv(template_path, index=False)
        logger.info("Ground-truth annotation template successfully generated at '%s'.", template_path)
    except Exception as e:
        logger.error("Error generating evaluation template: %s", e)

def main():
    parser = argparse.ArgumentParser(description="PII Redaction Tool for Word Documents")
    parser.add_argument(
        "--input", 
        type=str, 
        default="input/Red Herring Prospectus.docx", 
        help="Path to input DOCX file"
    )
    parser.add_argument(
        "--output", 
        type=str, 
        default="output/Red_Herring_Prospectus_Redacted.docx", 
        help="Path to output redacted DOCX file"
    )
    parser.add_argument(
        "--log", 
        type=str, 
        default="replacement_log.csv", 
        help="Path to save the replacement log CSV"
    )
    parser.add_argument(
        "--pdf-source", 
        type=str, 
        default="/Users/abhimansinghsaharan/.gemini/antigravity/brain/591a318f-d041-43f9-8eb7-96ae54c76d0c/media__1784295757960.pdf", 
        help="Path to PDF source if DOCX is missing"
    )
    
    args = parser.parse_args()
    setup_directories()

    input_docx = args.input
    output_docx = args.output
    log_csv = args.log
    pdf_source = args.pdf-source if hasattr(args, "pdf-source") else args.pdf_source

    # Step 1: Handle Missing File and convert PDF if available
    if not os.path.exists(input_docx):
        logger.warning("Input DOCX file '%s' not found.", input_docx)
        if os.path.exists(pdf_source):
            try:
                convert_pdf_to_docx(pdf_source, input_docx)
            except Exception as e:
                logger.critical("Failed to setup input document from PDF: %s", e)
                sys.exit(1)
        else:
            logger.critical(
                "Neither input DOCX '%s' nor source PDF '%s' could be found. Cannot proceed.", 
                input_docx, pdf_source
            )
            sys.exit(1)

    # Step 2: Verification of empty/invalid document before redacting
    try:
        from docx import Document
        doc = Document(input_docx)
        # Check if empty
        text_content = False
        for p in doc.paragraphs:
            if p.text.strip():
                text_content = True
                break
        if not text_content:
            for t in doc.tables:
                for row in t.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            if p.text.strip():
                                text_content = True
                                break
        if not text_content:
            logger.error("Input document '%s' is empty or contains no readable text.", input_docx)
            sys.exit(1)
    except Exception as e:
        logger.critical("Input DOCX '%s' is corrupted or invalid: %s", input_docx, e)
        sys.exit(1)

    # Step 3: Run PII Redaction
    logger.info("Initializing PII Redaction Engine...")
    try:
        redactor = DocxRedactor(use_presidio=True, use_spacy=True)
    except RuntimeError as e:
        logger.critical("Dependency or model error initializing detectors: %s", e)
        sys.exit(1)

    logger.info("Scanning and redacting document '%s'...", input_docx)
    try:
        total_redactions = redactor.redact_document(input_docx, output_docx)
        logger.info("Successfully completed redaction. Total PII occurrences replaced: %d", total_redactions)
        
        # Save replacement log
        redactor.save_replacement_log(log_csv)
        
        # Step 4: Generate template for NLP evaluation
        template_csv = "evaluation/ground_truth_template.csv"
        generate_evaluation_template(input_docx, template_csv)
        
        logger.info("PII Redaction finished successfully. Output: '%s'", output_docx)
        
    except Exception as e:
        logger.critical("Failed during document processing: %s", e)
        sys.exit(1)

if __name__ == "__main__":
    main()
