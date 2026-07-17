import os
import logging
from typing import List, Tuple
import pandas as pd
from docx import Document
from docx.text.paragraph import Paragraph

from src.detector import HybridDetector
from src.replacer import PIIReplacer
from src.utils import get_paragraph_runs

logger = logging.getLogger(__name__)

class DocxRedactor:
    """
    Orchestrates the detection and redaction of PII in Word (DOCX) documents.
    """
    def __init__(self, use_presidio: bool = True, use_spacy: bool = True):
        self.detector = HybridDetector(use_presidio=use_presidio, use_spacy=use_spacy)
        self.replacer = PIIReplacer()
        # Log of redactions: (original, replacement, entity_type)
        self.redaction_log: List[Tuple[str, str, str]] = []

    def redact_paragraph(self, paragraph: Paragraph) -> int:
        """
        Redact PII in a single paragraph, preserving inline formatting.
        Returns the number of redactions made.
        """
        runs = get_paragraph_runs(paragraph)
        if not runs:
            return 0
            
        full_text = "".join(run.text for run in runs)
        if not full_text.strip():
            return 0
            
        entities = self.detector.detect(full_text)
        if not entities:
            return 0

        # Sort entities in reverse order of start index to keep offsets valid for preceding entities
        entities_sorted = sorted(entities, key=lambda e: e.start, reverse=True)
        redactions_count = 0

        for entity in entities_sorted:
            start = entity.start
            end = entity.end
            original_text = entity.text
            entity_type = entity.entity_type
            
            replacement = self.replacer.get_replacement(original_text, entity_type)
            self.redaction_log.append((original_text, replacement, entity_type))
            redactions_count += 1
            
            # Re-map run ranges dynamically for the current state of runs
            run_ranges = []
            curr = 0
            for run in runs:
                text_len = len(run.text)
                run_ranges.append((curr, curr + text_len))
                curr += text_len
                
            # Find runs overlapping with [start, end)
            overlap_indices = [
                i for i, (rs, re) in enumerate(run_ranges)
                if rs < end and re > start
            ]
            
            if not overlap_indices:
                continue
                
            first_run_idx = overlap_indices[0]
            last_run_idx = overlap_indices[-1]
            
            # First run: keep prefix, add replacement
            rs_first, re_first = run_ranges[first_run_idx]
            prefix = runs[first_run_idx].text[:start - rs_first]
            
            # Last run: keep suffix
            rs_last, re_last = run_ranges[last_run_idx]
            suffix = runs[last_run_idx].text[end - rs_last:]
            
            if first_run_idx == last_run_idx:
                runs[first_run_idx].text = prefix + replacement + suffix
            else:
                runs[first_run_idx].text = prefix + replacement
                for idx in range(first_run_idx + 1, last_run_idx):
                    runs[idx].text = ""
                runs[last_run_idx].text = suffix
                
        return redactions_count

    def redact_table(self, table) -> int:
        """
        Recursively redact all cells and paragraphs inside a table.
        """
        redactions = 0
        for row in table.rows:
            for cell in row.cells:
                # Process paragraphs in the cell
                for p in cell.paragraphs:
                    redactions += self.redact_paragraph(p)
                # Recursively process nested tables inside the cell
                for t in cell.tables:
                    redactions += self.redact_table(t)
        return redactions

    def redact_document(self, input_path: str, output_path: str) -> int:
        """
        Load a DOCX document, redact all PII, and save the redacted version.
        """
        if not os.path.exists(input_path):
            raise FileNotFoundError(f"Input file not found: {input_path}")
            
        try:
            doc = Document(input_path)
        except Exception as e:
            logger.error("Failed to load DOCX: %s", e)
            raise ValueError(f"Failed to load or parse DOCX document: {e}") from e

        redactions = 0

        # 1. Redact body paragraphs
        for p in doc.paragraphs:
            redactions += self.redact_paragraph(p)

        # 2. Redact tables
        for t in doc.tables:
            redactions += self.redact_table(t)

        # 3. Redact headers and footers
        for section in doc.sections:
            if section.header:
                for p in section.header.paragraphs:
                    redactions += self.redact_paragraph(p)
                for t in section.header.tables:
                    redactions += self.redact_table(t)
            if section.footer:
                for p in section.footer.paragraphs:
                    redactions += self.redact_paragraph(p)
                for t in section.footer.tables:
                    redactions += self.redact_table(t)

        # 4. Redact shapes/textboxes in XML
        try:
            for txbx in doc.element.xpath('.//w:txbxContent'):
                for p_el in txbx.xpath('.//w:p'):
                    p = Paragraph(p_el, doc._body)
                    redactions += self.redact_paragraph(p)
        except Exception as e:
            logger.warning("Error redacting textboxes in XML: %s", e)

        # Save the redacted document
        try:
            # Ensure output directory exists
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            doc.save(output_path)
        except Exception as e:
            logger.error("Failed to save output DOCX: %s", e)
            raise IOError(f"Failed to save redacted document: {e}") from e

        return redactions

    def save_replacement_log(self, log_path: str) -> None:
        """
        Export the replacement log to a CSV file.
        """
        if not self.redaction_log:
            logger.warning("No redaction log entries to save.")
            return

        # Convert list of tuples to DataFrame and drop duplicates to keep the log clean and deduplicated
        df = pd.DataFrame(self.redaction_log, columns=["Original", "Replacement", "Entity Type"])
        df = df.drop_duplicates().reset_index(drop=True)
        
        try:
            df.to_csv(log_path, index=False)
            logger.info("Replacement log saved successfully to %s", log_path)
        except Exception as e:
            logger.error("Failed to save replacement log to CSV: %s", e)
            raise IOError(f"Failed to save replacement log CSV: {e}") from e
